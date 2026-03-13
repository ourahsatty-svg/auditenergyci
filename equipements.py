from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.database import get_db
from app import models
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, tempfile
from datetime import datetime

router = APIRouter()

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True
    return row

@router.get("/generer/{audit_id}")
def generer_rapport(audit_id: int, db: Session = Depends(get_db)):
    audit = db.query(models.Audit).filter(models.Audit.id == audit_id).first()
    if not audit:
        return {"error": "Audit non trouvé"}

    equipements = db.query(models.Equipement).filter(models.Equipement.audit_id == audit_id).all()
    consomm = db.query(models.ConsommationMensuelle).filter(
        models.ConsommationMensuelle.audit_id == audit_id
    ).order_by(models.ConsommationMensuelle.mois).all()
    apes = db.query(models.APE).filter(
        models.APE.audit_id == audit_id
    ).order_by(models.APE.priorite).all()

    total_kwh = sum(c.electricite_kwh + c.gaz_kwh + c.fioul_kwh for c in consomm)
    total_fcfa = sum(c.cout_fcfa for c in consomm)
    total_co2 = total_kwh * 0.47 / 1000
    potentiel_kwh = sum(a.economie_kwh_an for a in apes)
    potentiel_fcfa = sum(a.economie_fcfa_an for a in apes)

    doc = Document()

    # --- Style général ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- PAGE DE GARDE ---
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT D'AUDIT ÉNERGÉTIQUE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 110, 86)  # vert DGE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(audit.nom_entreprise.upper())
    run2.bold = True
    run2.font.size = Pt(16)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Conforme au modèle DGE — Direction Générale de l'Énergie\n")
    info.add_run(f"Normes : EN 16247 / ISO 50002\n")
    info.add_run(f"Localisation : {audit.localisation}\n")
    info.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y')}\n")
    info.add_run(f"Auditeur : {audit.auditeur}\n")
    doc.add_page_break()

    # --- SECTION 0 : Sigles ---
    set_heading(doc, "0. Liste des sigles et abréviations", 1, (15, 110, 86))
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Sigle"
    t.rows[0].cells[1].text = "Signification"
    sigles = [
        ("APE", "Action de Performance Énergétique"),
        ("CIE", "Compagnie Ivoirienne d'Électricité"),
        ("DGE", "Direction Générale de l'Énergie"),
        ("EN 16247", "Norme européenne audits énergétiques"),
        ("ISO 50002", "Norme internationale audits énergétiques"),
        ("IPMVP", "Protocole International de Mesure et Vérification"),
        ("ROI", "Return On Investment — Retour sur Investissement"),
        ("UES", "Usage Énergétique Significatif"),
        ("kWh", "Kilowattheure"),
        ("FCFA", "Franc CFA — monnaie en Côte d'Ivoire"),
        ("tCO₂", "Tonne équivalent CO₂"),
    ]
    for s in sigles:
        add_table_row(t, s)
    doc.add_page_break()

    # --- SECTION 1 : RÉSUMÉ EXÉCUTIF ---
    set_heading(doc, "1. Résumé exécutif", 1, (15, 110, 86))
    doc.add_paragraph(
        f"Le présent rapport présente les résultats de l'audit énergétique réalisé au sein de "
        f"{audit.nom_entreprise}, entreprise du secteur {audit.secteur or 'industriel'} "
        f"située à {audit.localisation}. L'audit a été conduit conformément aux normes EN 16247 "
        f"et ISO 50002, selon le modèle de cahier des charges de la Direction Générale de l'Énergie "
        f"de Côte d'Ivoire."
    )

    doc.add_paragraph("Principaux résultats :")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].text = "Indicateur"
    t2.rows[0].cells[1].text = "Valeur"
    resume_data = [
        ("Consommation totale annuelle", f"{total_kwh:,.0f} kWh"),
        ("Coût énergétique annuel", f"{total_fcfa:,.0f} FCFA"),
        ("Émissions CO₂ équivalent", f"{total_co2:,.1f} tCO₂/an"),
        ("Potentiel d'économie d'énergie", f"{potentiel_kwh:,.0f} kWh/an ({potentiel_kwh/total_kwh*100:.0f}% si >0)"),
        ("Potentiel d'économie financière", f"{potentiel_fcfa:,.0f} FCFA/an"),
        ("Nombre d'APE identifiées", str(len(apes))),
    ]
    for row in resume_data:
        add_table_row(t2, row, bold_first=True)
    doc.add_page_break()

    # --- SECTION 2 : PRÉSENTATION ENTREPRISE ---
    set_heading(doc, "2. Présentation de l'entreprise", 1, (15, 110, 86))
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = 'Table Grid'
    t3.rows[0].cells[0].text = "Information"
    t3.rows[0].cells[1].text = "Détail"
    info_data = [
        ("Raison sociale", audit.nom_entreprise),
        ("Secteur d'activité", audit.secteur or "N/A"),
        ("Type de site", audit.type_site),
        ("Localisation", audit.localisation or "N/A"),
        ("Surface", f"{audit.surface_m2:,.0f} m²" if audit.surface_m2 else "N/A"),
        ("Effectif", f"{audit.nb_employes} employés" if audit.nb_employes else "N/A"),
        ("Fournisseur énergie", audit.fournisseur_energie or "CIE"),
        ("Énergie principale", audit.type_energie or "Électricité"),
    ]
    for row in info_data:
        add_table_row(t3, row, bold_first=True)
    doc.add_page_break()

    # --- SECTION 3 : BILAN ENERGETIQUE ---
    set_heading(doc, "3. Bilan énergétique", 1, (15, 110, 86))
    set_heading(doc, "3.1. Consommations mensuelles", 2)

    if consomm:
        t4 = doc.add_table(rows=1, cols=4)
        t4.style = 'Table Grid'
        t4.rows[0].cells[0].text = "Mois"
        t4.rows[0].cells[1].text = "Électricité (kWh)"
        t4.rows[0].cells[2].text = "Gaz (kWh)"
        t4.rows[0].cells[3].text = "Coût (FCFA)"
        mois_noms = ["", "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
                     "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"]
        for c in consomm:
            add_table_row(t4, [
                mois_noms[c.mois] if c.mois <= 12 else str(c.mois),
                f"{c.electricite_kwh:,.0f}",
                f"{c.gaz_kwh:,.0f}",
                f"{c.cout_fcfa:,.0f}"
            ])
        total_row = t4.add_row()
        total_row.cells[0].text = "TOTAL"
        total_row.cells[1].text = f"{total_kwh:,.0f}"
        total_row.cells[2].text = ""
        total_row.cells[3].text = f"{total_fcfa:,.0f}"
        for cell in total_row.cells:
            cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None

    # --- SECTION 4 : EQUIPEMENTS ---
    set_heading(doc, "4. Description des installations", 1, (15, 110, 86))
    if equipements:
        t5 = doc.add_table(rows=1, cols=5)
        t5.style = 'Table Grid'
        headers = ["Équipement", "Puissance (kW)", "Heures/an", "Conso. (kWh/an)", "% total"]
        for i, h in enumerate(headers):
            t5.rows[0].cells[i].text = h
        for e in equipements:
            pct = (e.consommation_kwh_an / total_kwh * 100) if total_kwh > 0 else 0
            add_table_row(t5, [
                e.nom,
                str(e.puissance_kw),
                str(e.heures_an),
                f"{e.consommation_kwh_an:,.0f}" if e.consommation_kwh_an else "N/A",
                f"{pct:.1f}%"
            ])
    doc.add_page_break()

    # --- SECTION 5 : APE ---
    set_heading(doc, "5. Actions d'amélioration énergétique (APE)", 1, (15, 110, 86))
    doc.add_paragraph(
        "Les actions d'amélioration énergétique suivantes ont été identifiées lors de l'audit. "
        "Elles sont classées par ordre de priorité selon le retour sur investissement."
    )

    for i, ape in enumerate(apes, 1):
        set_heading(doc, f"APE {i} — {ape.titre}", 2)
        doc.add_paragraph(ape.description)
        t6 = doc.add_table(rows=5, cols=2)
        t6.style = 'Table Grid'
        ape_data = [
            ("Économie d'énergie estimée", f"{ape.economie_kwh_an:,.0f} kWh/an"),
            ("Économie financière", f"{ape.economie_fcfa_an:,.0f} FCFA/an"),
            ("Investissement requis", f"{ape.investissement_fcfa:,.0f} FCFA"),
            ("Temps de retour sur investissement", f"{ape.roi_mois:.1f} mois"),
            ("Réduction CO₂", f"{ape.reduction_co2_t_an:.2f} tCO₂/an"),
        ]
        for j, (k, v) in enumerate(ape_data):
            t6.rows[j].cells[0].text = k
            t6.rows[j].cells[1].text = v
            t6.rows[j].cells[0].paragraphs[0].runs[0].bold = True if t6.rows[j].cells[0].paragraphs[0].runs else None
        doc.add_paragraph()
    doc.add_page_break()

    # --- SECTION 6 : PLAN D'ACTION ---
    set_heading(doc, "6. Plan d'action et calendrier", 1, (15, 110, 86))
    if apes:
        t7 = doc.add_table(rows=1, cols=4)
        t7.style = 'Table Grid'
        t7.rows[0].cells[0].text = "Priorité"
        t7.rows[0].cells[1].text = "Action"
        t7.rows[0].cells[2].text = "Investissement (FCFA)"
        t7.rows[0].cells[3].text = "ROI (mois)"
        for ape in apes:
            add_table_row(t7, [
                f"P{ape.priorite}",
                ape.titre,
                f"{ape.investissement_fcfa:,.0f}",
                f"{ape.roi_mois:.1f}"
            ])
    doc.add_page_break()

    # --- SECTION 7 : CONCLUSION ---
    set_heading(doc, "7. Conclusion", 1, (15, 110, 86))
    doc.add_paragraph(
        f"L'audit énergétique de {audit.nom_entreprise} a permis d'identifier un potentiel "
        f"d'économie d'énergie de {potentiel_kwh:,.0f} kWh/an, représentant une économie "
        f"financière estimée à {potentiel_fcfa:,.0f} FCFA/an. "
        f"La mise en œuvre des {len(apes)} APE identifiées permettra également de réduire les "
        f"émissions de CO₂ de manière significative, contribuant ainsi aux objectifs nationaux "
        f"de transition énergétique de la Côte d'Ivoire."
    )
    doc.add_paragraph(
        "Ce rapport a été établi conformément au cahier des charges de la Direction Générale "
        "de l'Énergie (DGE) et aux normes EN 16247 / ISO 50002."
    )

    # --- SAUVEGARDE ---
    tmp_dir = tempfile.mkdtemp()
    filename = f"rapport_audit_{audit_id}_{audit.nom_entreprise.replace(' ', '_')}.docx"
    filepath = os.path.join(tmp_dir, filename)
    doc.save(filepath)

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
